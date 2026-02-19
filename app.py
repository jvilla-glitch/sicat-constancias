# -*- coding: utf-8 -*-
from flask import Flask, render_template, request, jsonify, send_file, redirect, url_for, flash
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from models import db, Constancia, Usuario
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from io import BytesIO

app = Flask(__name__, 
            static_folder='static',
            static_url_path='/static')

# ============================================
# CONFIGURACIÓN DE BASE DE DATOS
# ============================================
DATABASE_URL = os.environ.get('DATABASE_URL')

if DATABASE_URL:
    # Render proporciona postgres:// pero SQLAlchemy necesita postgresql://
    if DATABASE_URL.startswith("postgres://"):
        DATABASE_URL = DATABASE_URL.replace("postgres://", "postgresql://", 1)
    app.config['SQLALCHEMY_DATABASE_URI'] = DATABASE_URL
else:
    # Fallback a SQLite para desarrollo local
    app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///constancias.db'

app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# ============================================
# CONFIGURACIÓN DE UPLOADS - CRÍTICO PARA RENDER
# ============================================
# Detectar si estamos en Render
if os.environ.get('RENDER'):
    # En Render usar /tmp que tiene permisos de escritura
    app.config['UPLOAD_FOLDER'] = '/tmp/uploads'
    print("🌐 Modo PRODUCCIÓN (Render) - Usando /tmp/uploads")
else:
    # En local usar carpeta uploads normal
    app.config['UPLOAD_FOLDER'] = 'uploads'
    print("💻 Modo LOCAL - Usando ./uploads")

# Límite de tamaño de archivo: 16MB
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

# Obtener SECRET_KEY de variable de entorno o usar una por defecto
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'tu_clave_secreta_muy_segura_cambiarla_en_produccion')

db.init_app(app)

# Configurar Flask-Login
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'
login_manager.login_message = 'Por favor inicia sesión para acceder a esta página.'
login_manager.login_message_category = 'warning'

# Asegurar que exista la carpeta de uploads
try:
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    print(f"✅ Carpeta de uploads creada/verificada: {app.config['UPLOAD_FOLDER']}")
except Exception as e:
    print(f"⚠️ No se pudo crear carpeta de uploads: {e}")

@login_manager.user_loader
def load_user(user_id):
    return Usuario.query.get(int(user_id))

# Inicializar base de datos y crear usuarios predeterminados
def inicializar_usuarios():
    """Crea los usuarios iniciales si no existen"""
    usuarios_iniciales = [
        'ALEJANDRA SEPULVEDA',
        'JOSUE VILLA',
        'BRENDA GONZALEZ',
        'ERIK CHAVEZ',
        'MARLEN ROCHA',
        'LILIA PEÑA'
    ]
    
    for nombre in usuarios_iniciales:
        usuario_existe = Usuario.query.filter_by(nombre=nombre).first()
        if not usuario_existe:
            nuevo_usuario = Usuario(nombre=nombre)
            nuevo_usuario.set_password('cis2025')  # Contraseña predeterminada
            db.session.add(nuevo_usuario)
    
    db.session.commit()
    print("✅ Usuarios inicializados correctamente")

with app.app_context():
    db.create_all()
    inicializar_usuarios()

# ============================================
# RUTAS PRINCIPALES
# ============================================

@app.route('/')
@login_required
def index():
    """Ruta raíz redirige al dashboard"""
    return redirect(url_for('dashboard'))

@app.route('/dashboard')
@login_required
def dashboard():
    """Página principal del dashboard"""
    return render_template('dashboard.html', usuario=current_user)

# ============================================
# RUTAS DE AUTENTICACIÓN
# ============================================

@app.route('/login', methods=['GET', 'POST'])
def login():
    """Página de inicio de sesión"""
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        nombre = request.form.get('nombre')
        password = request.form.get('password')
        recordar = request.form.get('recordar') == 'on'
        
        usuario = Usuario.query.filter_by(nombre=nombre).first()
        
        if usuario and usuario.check_password(password):
            if not usuario.activo:
                flash('Tu cuenta ha sido desactivada. Contacta al administrador.', 'danger')
                return redirect(url_for('login'))
            
            # Actualizar último acceso
            usuario.ultimo_acceso = datetime.utcnow()
            db.session.commit()
            
            login_user(usuario, remember=recordar)
            flash(f'¡Bienvenido(a) {usuario.nombre}!', 'success')
            
            # Redirigir a la página solicitada o al dashboard
            next_page = request.args.get('next')
            return redirect(next_page) if next_page else redirect(url_for('dashboard'))
        else:
            flash('Nombre de usuario o contraseña incorrectos', 'danger')
    
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    """Cerrar sesión"""
    logout_user()
    flash('Has cerrado sesión exitosamente', 'info')
    return redirect(url_for('login'))

# ============================================
# RUTAS DE MANUALES
# ============================================

@app.route('/manual')
@app.route('/manuales')
@login_required
def listar_manuales():
    """Página de listado de constancias manuales"""
    return render_template('manual/index.html', usuario=current_user)

@app.route('/manual/nueva')
@login_required
def nueva_manual():
    """Página para crear nueva constancia manual"""
    return render_template('manual/formulario.html', usuario=current_user, modo='crear')

@app.route('/manual/editar/<int:id>')
@login_required
def editar_manual(id):
    """Página para editar constancia manual existente"""
    constancia = Constancia.query.get_or_404(id)
    return render_template('manual/formulario.html', usuario=current_user, constancia=constancia, modo='editar')

# ============================================
# RUTAS DE MASIVAS
# ============================================

@app.route('/masivo')
@app.route('/masivas')
@login_required
def pagina_masivas():
    """Página de importación masiva de constancias"""
    return render_template('masivo/index.html', usuario=current_user)

@app.route('/masivo/importar', methods=['POST'])
@login_required
def importar_masivo():
    """Importar constancias desde archivo Excel/CSV - PROCESAMIENTO EN MEMORIA"""
    try:
        # Verificar que se envió un archivo
        if 'file' not in request.files:
            return jsonify({
                'success': False, 
                'error': 'No se envió ningún archivo. Por favor selecciona un archivo Excel o CSV.'
            }), 400
        
        file = request.files['file']
        
        # Verificar que se seleccionó un archivo
        if file.filename == '':
            return jsonify({
                'success': False, 
                'error': 'No se seleccionó ningún archivo. Por favor selecciona un archivo.'
            }), 400
        
        # Verificar extensión del archivo
        if not (file.filename.endswith('.xlsx') or file.filename.endswith('.xls') or file.filename.endswith('.csv')):
            return jsonify({
                'success': False, 
                'error': 'Formato de archivo no soportado. Use archivos .xlsx, .xls o .csv'
            }), 400
        
        # ============================================
        # PROCESAR ARCHIVO DIRECTAMENTE EN MEMORIA
        # No se guarda en disco - Compatible con Render
        # ============================================
        try:
            if file.filename.endswith('.csv'):
                df = pd.read_csv(file)
            else:
                df = pd.read_excel(file)
        except Exception as e:
            return jsonify({
                'success': False,
                'error': f'Error al leer el archivo: {str(e)}. Verifica que el archivo no esté corrupto.'
            }), 400
        
        # Verificar que el archivo no esté vacío
        if df.empty:
            return jsonify({
                'success': False,
                'error': 'El archivo está vacío. Por favor verifica el contenido del archivo.'
            }), 400
        
        # Verificar columnas requeridas
        columnas_requeridas = [
            'folio', 'fecha', 'vigencia', 'expediente', 'nombre_propietario',
            'domicilio_propietario', 'giro', 'denominado', 'ubicado', 'entre_calles',
            'colonia', 'ciudad', 'codigo_postal', 'nombre_comisionado'
        ]
        
        columnas_faltantes = [col for col in columnas_requeridas if col not in df.columns]
        if columnas_faltantes:
            return jsonify({
                'success': False,
                'error': f'Faltan las siguientes columnas requeridas: {", ".join(columnas_faltantes)}'
            }), 400
        
        constancias_creadas = []
        errores = []
        
        # Procesar cada fila del archivo
        for index, row in df.iterrows():
            try:
                # Verificar que el folio no exista
                folio = str(row['folio']).strip()
                
                # Saltar filas con folio vacío
                if not folio or folio == 'nan' or folio == '':
                    continue
                
                folio_existente = Constancia.query.filter_by(folio=folio).first()
                
                if folio_existente:
                    errores.append(f"Fila {index + 2}: El folio {folio} ya existe")
                    continue
                
                # Función auxiliar para limpiar valores
                def limpiar_valor(valor):
                    if pd.isna(valor) or str(valor).strip() == '' or str(valor) == 'nan':
                        return ''
                    return str(valor).strip()
                
                # Crear la constancia
                constancia = Constancia(
                    fecha=limpiar_valor(row['fecha']),
                    folio=folio,
                    vigencia=limpiar_valor(row['vigencia']),
                    expediente=limpiar_valor(row['expediente']),
                    nombre_propietario=limpiar_valor(row['nombre_propietario']),
                    domicilio_propietario=limpiar_valor(row['domicilio_propietario']),
                    giro=limpiar_valor(row['giro']),
                    denominado=limpiar_valor(row['denominado']),
                    ubicado=limpiar_valor(row['ubicado']),
                    entre_calles=limpiar_valor(row['entre_calles']),
                    colonia=limpiar_valor(row['colonia']),
                    ciudad=limpiar_valor(row['ciudad']),
                    codigo_postal=limpiar_valor(row['codigo_postal']),
                    nombre_comisionado=limpiar_valor(row['nombre_comisionado']),
                    recibo_pago=limpiar_valor(row.get('recibo_pago', '')),
                    referencia_comprobante=limpiar_valor(row.get('referencia_comprobante', '')),
                    constancias_avala=limpiar_valor(row.get('constancias_avala', ''))
                )
                
                db.session.add(constancia)
                constancias_creadas.append(constancia.to_dict())
                
            except Exception as e:
                errores.append(f"Fila {index + 2}: {str(e)}")
        
        # Guardar todas las constancias
        if constancias_creadas:
            db.session.commit()
            mensaje = f'✅ Se importaron {len(constancias_creadas)} constancias exitosamente'
        else:
            mensaje = '⚠️ No se importó ninguna constancia'
        
        if errores:
            mensaje += f' | ⚠️ {len(errores)} errores encontrados'
        
        return jsonify({
            'success': True,
            'mensaje': mensaje,
            'total_importadas': len(constancias_creadas),
            'total_errores': len(errores),
            'constancias_creadas': constancias_creadas,
            'errores': errores
        })
        
    except Exception as e:
        db.session.rollback()
        print(f"❌ Error en importación masiva: {str(e)}")
        return jsonify({
            'success': False, 
            'error': f'Error al procesar el archivo: {str(e)}'
        }), 500

# ============================================
# FUNCIÓN PARA GENERAR DOCUMENTO WORD
# ============================================

def crear_documento_word(constancia):
    """Genera un documento Word con los datos de la constancia"""
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    def configurar_formato_parrafo(p, alineacion=None):
        """Configura el formato estándar de párrafo: espaciado 0, interlineado sencillo"""
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        if alineacion:
            p.alignment = alineacion
    
    # ENCABEZADO - Alineado a la derecha
    p = doc.add_paragraph()
    configurar_formato_parrafo(p, WD_ALIGN_PARAGRAPH.RIGHT)
    
    run = p.add_run(f'FOLIO No.: {constancia.folio}-2026\n')
    run.bold = True
    run.font.size = Pt(12)
    
    run = p.add_run('ASUNTO: CONSTANCIA SANITARIA\n')
    run.bold = True
    run.font.size = Pt(12)
    
    run = p.add_run(f'VIGENCIA: {constancia.vigencia}\n')
    run.bold = True
    run.font.size = Pt(12)
    
    run = p.add_run(f'EXPEDIENTE: {constancia.expediente}\n')
    run.bold = True
    run = p.add_run(constancia.fecha)
    run.font.size = Pt(10)
    
    # Párrafo vacío (un solo enter)
    p = doc.add_paragraph()
    configurar_formato_parrafo(p)

    p = doc.add_paragraph()
    configurar_formato_parrafo(p)
    
    # DESTINATARIO
    p = doc.add_paragraph()
    configurar_formato_parrafo(p)
    
    run = p.add_run(constancia.nombre_propietario + '\n')
    run.bold = True
    run.font.size = Pt(14)
    
    run = p.add_run(constancia.domicilio_propietario)
    run.bold = True
    run.font.size = Pt(14)
    
    # Párrafo vacío (un solo enter)
    p = doc.add_paragraph()
    configurar_formato_parrafo(p)
    
    # CUERPO - Párrafo 1
    p = doc.add_paragraph()
    configurar_formato_parrafo(p, WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    run = p.add_run('Se expide la presente ')
    run.font.size = Pt(12)
    
    run = p.add_run('CONSTANCIA SANITARIA')
    run.bold = True
    run.font.size = Pt(12)
    
    run = p.add_run(' a petición del interesado en virtud de la solicitud mediante la cual manifiesta, '
                    'bajo protesta de decir verdad, que el establecimiento reúne los requisitos sanitarios '
                    'para su funcionamiento, destinado a:')
    run.font.size = Pt(12)
    
    # Párrafo vacío (un solo enter)
    p = doc.add_paragraph()
    configurar_formato_parrafo(p)
    
    # DATOS DEL ESTABLECIMIENTO
    p = doc.add_paragraph()
    configurar_formato_parrafo(p)
    p.paragraph_format.left_indent = Inches(0.5)
    
    run = p.add_run(constancia.giro + '\n')
    run.bold = True
    run.font.size = Pt(12)
    
    run = p.add_run(f'DENOMINADO: {constancia.denominado}\n')
    run.font.size = Pt(12)
    
    run = p.add_run(f'UBICADO EN: {constancia.ubicado}\n')
    run.font.size = Pt(12)
    
    run = p.add_run(f'ENTRE LAS CALLES: {constancia.entre_calles}\n')
    run.font.size = Pt(12)
    
    run = p.add_run(f'COLONIA: {constancia.colonia}    CÓDIGO POSTAL: {constancia.codigo_postal}\n')
    run.font.size = Pt(12)
    
    run = p.add_run(constancia.ciudad)
    run.font.size = Pt(12)
    
    # Párrafo vacío (un solo enter)
    p = doc.add_paragraph()
    configurar_formato_parrafo(p)
    
    # CUERPO - Párrafo 2
    p = doc.add_paragraph()
    configurar_formato_parrafo(p, WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    run = p.add_run(
        'Lo anterior, en apego a lo contemplado en el artículo 21, fracción IV de la Ley sobre '
        'Operación y Funcionamiento de Establecimientos Destinados a la Producción, Distribución '
        'Venta y Consumo de Bebidas Alcohólicas del Estado de Sinaloa, 16, fracción III de su '
        'Reglamento, en relación con el artículo 372 de la Ley General de Salud, los artículos '
        '30, 32, 33 de su Reglamento de Control Sanitario de Productos y Servicios y, en '
        'consecuencia, con los artículos 8, 15, 16, 17, 28, 36 y 41 de la Ley General para el '
        'Control de Tabaco; con el apercibimiento que esta autoridad podrá hacer uso de la '
        'figura de la '
    )
    run.font.size = Pt(12)
    
    run = p.add_run('REVOCACIÓN')
    run.bold = True
    run.font.size = Pt(12)
    
    run = p.add_run(
        ' de este tipo de constancias sanitarias así requeridas cuando hubiere desobediencia en '
        'acatar las disposiciones sanitarias en los términos de la Ley y demás disposiciones '
        'legales aplicables.'
    )
    run.font.size = Pt(12)
    
    # Párrafo vacío (un solo enter)
    p = doc.add_paragraph()
    configurar_formato_parrafo(p)
    
    # Párrafo vacío (un solo enter)
    p = doc.add_paragraph()
    configurar_formato_parrafo(p)
    
    # Párrafo vacío (un solo enter)
    p = doc.add_paragraph()
    configurar_formato_parrafo(p)
    
    # FIRMA
    p = doc.add_paragraph()
    configurar_formato_parrafo(p, WD_ALIGN_PARAGRAPH.CENTER)
    
    run = p.add_run('A T E N T A M E N T E')
    run.bold = True
    run.font.size = Pt(12)
    
    # Espacios para firma
    for _ in range(3):
        p = doc.add_paragraph()
        configurar_formato_parrafo(p)
    
    # Cargo
    p = doc.add_paragraph()
    configurar_formato_parrafo(p, WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run('COMISIONADA ESTATAL PARA LA PROTECCION CONTRA RIESGOS SANITARIOS DE SINALOA')
    run.bold = True
    run.font.size = Pt(12)
    
    # Espacios para nombre
    for _ in range(2):
        p = doc.add_paragraph()
        configurar_formato_parrafo(p)
    
    # Nombre del comisionado
    p = doc.add_paragraph()
    configurar_formato_parrafo(p, WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(constancia.nombre_comisionado)
    run.bold = True
    run.font.size = Pt(12)
    
    return doc

@app.route('/generar-masivo', methods=['POST'])
@app.route('/masivo/generar-pdfs', methods=['POST'])
@login_required
def generar_masivo():
    """Generar documento Word con múltiples constancias"""
    try:
        data = request.get_json()
        ids = data.get('ids', [])
        
        if not ids:
            return jsonify({'error': 'No se especificaron IDs de constancias'}), 400
        
        doc_principal = None
        
        for id_constancia in ids:
            constancia = Constancia.query.get(id_constancia)
            if constancia:
                if doc_principal is None:
                    # Primera constancia - crear documento
                    doc_principal = crear_documento_word(constancia)
                else:
                    # Constancias subsecuentes - agregar salto de página y luego el contenido
                    doc_principal.add_page_break()
                    
                    # Agregar contenido de la siguiente constancia
                    doc_temp = crear_documento_word(constancia)
                    
                    # Copiar solo los párrafos (no elementos vacíos)
                    for paragraph in doc_temp.paragraphs:
                        new_para = doc_principal.add_paragraph()
                        new_para.paragraph_format.space_after = paragraph.paragraph_format.space_after
                        new_para.paragraph_format.space_before = paragraph.paragraph_format.space_before
                        new_para.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing
                        new_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
                        new_para.paragraph_format.left_indent = paragraph.paragraph_format.left_indent
                        
                        for run in paragraph.runs:
                            new_run = new_para.add_run(run.text)
                            new_run.bold = run.bold
                            new_run.font.size = run.font.size
        
        if doc_principal is None:
            return jsonify({'error': 'No se encontraron constancias válidas'}), 404
        
        # Guardar en buffer
        doc_buffer = BytesIO()
        doc_principal.save(doc_buffer)
        doc_buffer.seek(0)
        
        return send_file(
            doc_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'constancias_masivas_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        )
        
    except Exception as e:
        return jsonify({'error': f'Error al generar documentos: {str(e)}'}), 500

# ============================================
# API REST - CRUD DE CONSTANCIAS
# ============================================

@app.route('/api/constancias', methods=['GET'])
@login_required
def listar_constancias():
    """Obtener todas las constancias"""
    constancias = Constancia.query.order_by(Constancia.fecha_creacion.desc()).all()
    return jsonify([c.to_dict() for c in constancias])

@app.route('/api/constancias/<int:id>', methods=['GET'])
@login_required
def obtener_constancia(id):
    """Obtener una constancia por ID"""
    constancia = Constancia.query.get_or_404(id)
    return jsonify(constancia.to_dict())

@app.route('/api/constancias', methods=['POST'])
@login_required
def crear_constancia():
    """Crear una nueva constancia"""
    data = request.get_json()
    
    # Verificar que el folio no exista
    folio_existente = Constancia.query.filter_by(folio=data['folio']).first()
    if folio_existente:
        return jsonify({'error': 'El folio ya existe'}), 400
    
    try:
        constancia = Constancia(
            fecha=data['fecha'],
            folio=data['folio'],
            vigencia=data['vigencia'],
            expediente=data['expediente'],
            nombre_propietario=data['nombre_propietario'],
            domicilio_propietario=data['domicilio_propietario'],
            giro=data['giro'],
            denominado=data['denominado'],
            ubicado=data['ubicado'],
            entre_calles=data['entre_calles'],
            colonia=data['colonia'],
            ciudad=data['ciudad'],
            codigo_postal=data['codigo_postal'],
            nombre_comisionado=data['nombre_comisionado'],
            recibo_pago=data.get('recibo_pago'),
            referencia_comprobante=data.get('referencia_comprobante'),
            constancias_avala=data.get('constancias_avala')
        )
        
        db.session.add(constancia)
        db.session.commit()
        
        return jsonify(constancia.to_dict()), 201
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/constancias/<int:id>', methods=['PUT'])
@login_required
def actualizar_constancia(id):
    """Actualizar una constancia existente"""
    constancia = Constancia.query.get_or_404(id)
    data = request.get_json()
    
    # Verificar que el folio no esté duplicado (excepto el actual)
    if data['folio'] != constancia.folio:
        folio_existente = Constancia.query.filter_by(folio=data['folio']).first()
        if folio_existente:
            return jsonify({'error': 'El folio ya existe'}), 400
    
    try:
        constancia.fecha = data['fecha']
        constancia.folio = data['folio']
        constancia.vigencia = data['vigencia']
        constancia.expediente = data['expediente']
        constancia.nombre_propietario = data['nombre_propietario']
        constancia.domicilio_propietario = data['domicilio_propietario']
        constancia.giro = data['giro']
        constancia.denominado = data['denominado']
        constancia.ubicado = data['ubicado']
        constancia.entre_calles = data['entre_calles']
        constancia.colonia = data['colonia']
        constancia.ciudad = data['ciudad']
        constancia.codigo_postal = data['codigo_postal']
        constancia.nombre_comisionado = data['nombre_comisionado']
        constancia.recibo_pago = data.get('recibo_pago')
        constancia.referencia_comprobante = data.get('referencia_comprobante')
        constancia.constancias_avala = data.get('constancias_avala')
        
        db.session.commit()
        
        return jsonify(constancia.to_dict())
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/constancias/<int:id>', methods=['DELETE'])
@login_required
def eliminar_constancia(id):
    """Eliminar una constancia"""
    constancia = Constancia.query.get_or_404(id)
    
    try:
        db.session.delete(constancia)
        db.session.commit()
        return jsonify({'mensaje': 'Constancia eliminada exitosamente'})
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/constancia/<int:id>/pdf')
@login_required
def generar_pdf(id):
    """Generar documento Word de una constancia"""
    constancia = Constancia.query.get_or_404(id)
    
    try:
        # Generar documento Word
        doc = crear_documento_word(constancia)
        
        # Guardar en buffer
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        response = send_file(
            doc_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'constancia_{constancia.folio}.docx'
        )
        
        return response
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/estadisticas')
@login_required
def api_estadisticas():
    """Obtener estadísticas del dashboard"""
    try:
        total = Constancia.query.count()
        
        # Constancias de hoy
        hoy = datetime.now().date()
        constancias_hoy = Constancia.query.filter(
            db.func.date(Constancia.fecha_creacion) == hoy
        ).count()
        
        # Constancias de esta semana
        hace_una_semana = datetime.now() - pd.Timedelta(days=7)
        constancias_semana = Constancia.query.filter(
            Constancia.fecha_creacion >= hace_una_semana
        ).count()
        
        # Constancias de este mes
        primer_dia_mes = datetime.now().replace(day=1)
        constancias_mes = Constancia.query.filter(
            Constancia.fecha_creacion >= primer_dia_mes
        ).count()
        
        return jsonify({
            'total': total,
            'hoy': constancias_hoy,
            'semana': constancias_semana,
            'mes': constancias_mes
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500
@app.route('/exportar-datos')
@login_required
def exportar_datos():
    import csv
    from io import StringIO
    constancias = Constancia.query.all()
    output = StringIO()
    cols = ['fecha','folio','vigencia','expediente','nombre_propietario',
            'domicilio_propietario','giro','denominado','ubicado','entre_calles',
            'colonia','ciudad','codigo_postal','nombre_comisionado',
            'recibo_pago','referencia_comprobante','constancias_avala']
    writer = csv.DictWriter(output, fieldnames=cols, extrasaction='ignore')
    writer.writeheader()
    writer.writerows([c.to_dict() for c in constancias])
    output.seek(0)
    return send_file(
        BytesIO(output.getvalue().encode('utf-8')),
        mimetype='text/csv',
        as_attachment=True,
        download_name='constancias_export.csv'
    )
if __name__ == '__main__':
    app.run(debug=True)












